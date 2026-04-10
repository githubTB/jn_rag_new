import logging
import re
from pathlib import Path

from docx import Document as DocxDocument
from docx.oxml.ns import qn
from docx.text.run import Run

from .base import BaseExtractor
from models.document import Document

logger = logging.getLogger(__name__)


class WordExtractor(BaseExtractor):
    """
    Extract text and tables from .docx files.
    Images are represented as [IMAGE] placeholders.
    """

    def __init__(self, file_path: str):
        self._file_path = file_path

    def extract(self) -> list[Document]:
        content = self._parse_docx(self._file_path)
        return [Document(page_content=content, metadata={"source": self._file_path})]

    # ------------------------------------------------------------------ #
    #  Internal parsing                                                    #
    # ------------------------------------------------------------------ #

    def _parse_docx(self, path: str) -> str:
        doc = DocxDocument(path)
        parts: list[str] = []
        paragraphs = list(doc.paragraphs)
        tables = list(doc.tables)

        for element in doc.element.body:
            tag = getattr(element, "tag", None)
            if not isinstance(tag, str):
                continue
            if tag.endswith("p") and paragraphs:
                para = paragraphs.pop(0)
                text = self._parse_paragraph(para, doc)
                parts.append(text.strip() if text.strip() else "")
            elif tag.endswith("tbl") and tables:
                table = tables.pop(0)
                parts.append(self._table_to_markdown(table))

        return "\n".join(p for p in parts if p)

    # ---- paragraph ---------------------------------------------------- #

    def _parse_paragraph(self, paragraph, doc: DocxDocument) -> str:
        content: list[str] = []
        hyperlink_url: str | None = None
        hyperlink_text_parts: list[str] = []
        collecting_field_text = False

        for child in paragraph._element:
            tag = child.tag
            if tag == qn("w:r"):
                run = Run(child, paragraph)
                # Legacy HYPERLINK field handling
                fld_chars = child.findall(qn("w:fldChar"))
                instr_texts = child.findall(qn("w:instrText"))
                if fld_chars or instr_texts:
                    for instr in instr_texts:
                        if instr.text and "HYPERLINK" in instr.text:
                            m = re.search(r'HYPERLINK\s+"([^"]+)"', instr.text, re.IGNORECASE)
                            if m:
                                hyperlink_url = m.group(1)
                    for fc in fld_chars:
                        fc_type = fc.get(qn("w:fldCharType"))
                        if fc_type == "begin":
                            hyperlink_url = None
                            hyperlink_text_parts = []
                            collecting_field_text = False
                        elif fc_type == "separate" and hyperlink_url:
                            collecting_field_text = True
                        elif fc_type == "end":
                            if collecting_field_text and hyperlink_url:
                                display = "".join(hyperlink_text_parts).strip()
                                if display:
                                    content.append(f"[{display}]({hyperlink_url})")
                            hyperlink_url = None
                            hyperlink_text_parts = []
                            collecting_field_text = False

                target = hyperlink_text_parts if collecting_field_text else content
                self._process_run(run, doc, target)

            elif tag == qn("w:hyperlink"):
                r_id = child.get(qn("r:id"))
                link_texts = []
                for run_elem in child.findall(qn("w:r")):
                    r = Run(run_elem, paragraph)
                    if r.text:
                        link_texts.append(r.text)
                link_text = "".join(link_texts).strip()
                if r_id:
                    try:
                        rel = paragraph.part.rels.get(r_id)
                        if rel and rel.is_external:
                            link_text = f"[{link_text or rel.target_ref}]({rel.target_ref})"
                    except Exception:
                        pass
                if link_text:
                    content.append(link_text)

        return "".join(content)

    def _process_run(self, run: Run, doc: DocxDocument, buf: list[str]) -> None:
        """Append run text or [IMAGE] placeholder to buf."""
        # Use findall + Clark notation — python-docx xpath() has no namespaces= kwarg
        _A_BLIP = "{http://schemas.openxmlformats.org/drawingml/2006/main}blip"
        _R_EMBED = "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
        blips = run.element.findall(f".//{_A_BLIP}")
        has_image = False
        for blip in blips:
            if blip.get(_R_EMBED):
                buf.append("[IMAGE]")
                has_image = True
        if not has_image and run.text:
            buf.append(run.text)

    # ---- table -------------------------------------------------------- #

    def _table_to_markdown(self, table) -> str:
        if not table.rows:
            return ""
        total_cols = max(len(row.cells) for row in table.rows)
        rows_md: list[str] = []

        for i, row in enumerate(table.rows):
            cells = self._parse_row_cells(row, total_cols)
            rows_md.append("| " + " | ".join(cells) + " |")
            if i == 0:
                rows_md.append("| " + " | ".join(["---"] * total_cols) + " |")

        return "\n".join(rows_md)

    def _parse_row_cells(self, row, total_cols: int) -> list[str]:
        result = [""] * total_cols
        col = 0
        for cell in row.cells:
            while col < total_cols and result[col]:
                col += 1
            if col >= total_cols:
                break
            text = " ".join(p.text.strip() for p in cell.paragraphs if p.text.strip())
            span = getattr(cell, "grid_span", 1) or 1
            for i in range(span):
                if col + i < total_cols:
                    result[col + i] = text if i == 0 else ""
            col += span
        return result
